import streamlit as st
from PIL import Image
import os
import zipfile
import io
import requests
from docx import Document

# Configura Streamlit
st.set_page_config(page_title="Evaluador de Hoja de Vida ANEIAP", layout="wide")

# Imagen para la interfaz
imagen_aneiap = 'Evaluador Hoja de Vida ANEIAP UNINORTE.jpg'

with st.container():
    st.subheader("¿Qué tan listo estás para asumir un cargo de junta directiva Capitular? Descúbrelo aquí :lion:")
    st.title("Evaluador de Hoja de Vida ANEIAP UNINORTE")
    st.write("Con solo tu hoja de vida ANEIAP (en formato .docx) podrás averiguar qué tan preparado te encuentras para asumir un cargo dentro de la JDC-IC-CCP.")
    st.image(imagen_aneiap, use_container_width=True)

# Función para cargar documentos de funciones y perfil
def load_job_docs(cargo):
    base_path = "/content/drive/MyDrive/HERRAMIENTAS EN COLLAB/EVALUADOR HOJA DE VIDA ANEIAP/"
    job_docs = {
        "PC": {"funciones": base_path + "CARGOS JUNTA/FPC.docx", "perfil": base_path + "CARGOS JUNTA/PPC.docx"},
        "DCA": {"funciones": base_path + "CARGOS JUNTA/FDCA.docx", "perfil": base_path + "CARGOS JUNTA/PDCA.docx"},
        "DCC": {"funciones": base_path + "CARGOS JUNTA/FDCC.docx", "perfil": base_path + "CARGOS JUNTA/PDCC.docx"},
        "DCD": {"funciones": base_path + "CARGOS JUNTA/FDCD.docx", "perfil": base_path + "CARGOS JUNTA/PDCD.docx"},
        "DCF": {"funciones": base_path + "CARGOS JUNTA/FDCF.docx", "perfil": base_path + "CARGOS JUNTA/PDCF.docx"},
        "DCM": {"funciones": base_path + "CARGOS JUNTA/FDCM.docx", "perfil": base_path + "CARGOS JUNTA/PDCM.docx"},
        "CCP": {"funciones": base_path + "CARGOS JUNTA/FCCP.docx", "perfil": base_path + "CARGOS JUNTA/PCCP.docx"},
        "IC": {"funciones": base_path + "CARGOS JUNTA/FIC.docx", "perfil": base_path + "CARGOS JUNTA/PIC.docx"}
    }
    return job_docs[cargo]

# Función para comparar similitudes usando Llama3
def llama3_similarity(text, reference_text, api_key):
    """Usa la API de Llama3 para calcular la similitud entre dos textos."""
    url = "https://api.llama3.com/v1/similarity"
    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json"
    }
    payload = {
        "text1": text,
        "text2": reference_text
    }
    response = requests.post(url, json=payload, headers=headers)
    if response.status_code == 200:
        similarity_score = response.json()["similarity"]
        return similarity_score
    else:
        st.error("Error en la API de Llama3")
        return 0

# Extracción del contenido de interés (Experiencia ANEIAP)
def extract_experience_aneiap(doc):
    """Extrae el texto de la sección 'EXPERIENCIA EN ANEIAP'."""
    text = ""
    paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
    extract = False
    for para in paragraphs:
        if "EXPERIENCIA EN ANEIAP" in para.upper():
            extract = True
        elif extract and any(keyword in para.upper() for keyword in ["EVENTOS ORGANIZADOS", "ASISTENCIA A EVENTOS ANEIAP"]):
            break
        elif extract:
            text += para + "\n"
    return text.strip()

# Función para extraer documentos .docx de un archivo zip
def extract_docx_from_zip(uploaded_zip):
    """Extrae los archivos .docx de un archivo zip cargado."""
    docx_files = []
    with zipfile.ZipFile(uploaded_zip, "r") as zip_ref:
        for file in zip_ref.namelist():
            if file.endswith(".docx"):
                with zip_ref.open(file) as doc_file:
                    docx_files.append(io.BytesIO(doc_file.read()))
    return docx_files

# Generación del reporte
def generate_report(experience_text, func_text, profile_text, cargo, candidate_name, api_key):
    """Genera un reporte de análisis y lo exporta como archivo .docx."""
    # Separar los ítems de experiencia en ANEIAP
    experience_items = [item.strip() for item in experience_text.split('\n') if item.strip()]

    # Crear documento
    doc = Document()

    # Título del reporte
    title = doc.add_paragraph("REPORTE ANÁLISIS HOJA DE VIDA")
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run_title = title.runs[0]
    run_title.font.bold = True
    run_title.font.size = Pt(16)
    run_title.font.name = "Century Gothic"

    # Evaluar cada ítem de la sección "EXPERIENCIA EN ANEIAP"
    item_similarities_func = []
    item_similarities_prof = []

    for item in experience_items:
        # Usar Llama3 para comparar cada ítem con las funciones y perfil
        similarity_func = llama3_similarity(item, func_text, api_key)
        similarity_prof = llama3_similarity(item, profile_text, api_key)
        item_similarities_func.append(max(similarity_func, 0.0))
        item_similarities_prof.append(max(similarity_prof, 0.0))

    # Promediar las similitudes para las funciones y el perfil
    avg_similarity_func = sum(item_similarities_func) / len(item_similarities_func) if item_similarities_func else 0
    avg_similarity_prof = sum(item_similarities_prof) / len(item_similarities_prof) if item_similarities_prof else 0

    # Incluir los resultados en el reporte
    doc.add_paragraph(f"Similitud promedio con las funciones: {avg_similarity_func * 100:.2f}%")
    doc.add_paragraph(f"Similitud promedio con el perfil: {avg_similarity_prof * 100:.2f}%")

    # Conclusión
    if avg_similarity_func < 0.5 or avg_similarity_prof < 0.5:
        doc.add_paragraph(f"- Baja Concordancia (< 0.50): El análisis indica que {candidate_name} tiene una baja concordancia con los requisitos del cargo "
            f"de {cargo} y el perfil buscado. Esto sugiere que aunque el aspirante posee algunas experiencias relevantes, su historial actual "
            f"no cubre adecuadamente las competencias y responsabilidades necesarias para este rol crucial en la prevalencia del Capítulo. "
            f"Se aconseja a {candidate_name} enfocarse en mejorar su perfil profesional y desarrollar las habilidades necesarias para el cargo. "
            f"Este enfoque permitirá a {candidate_name} alinear mejor su perfil con los requisitos del puesto en futuras oportunidades.")
    elif avg_similarity_func < 0.75 or avg_similarity_prof < 0.75:
        doc.add_paragraph(f"- Buena Concordancia (≥ 0.50): El análisis muestra que {candidate_name} tiene una buena correspondencia con las funciones "
            f"del cargo de {cargo} y el perfil deseado. Aunque su experiencia en la asociación es relevante, existe margen para mejorar. "
            f"{candidate_name} muestra potencial para cumplir con el rol crucial en la prevalencia del Capítulo, pero se recomienda que continúe "
            f"desarrollando sus habilidades y acumulando más experiencia relacionada con el cargo objetivo. Su candidatura debe ser considerada "
            f"con la recomendación de enriquecimiento adicional.")
    else:
        doc.add_paragraph(f"- Alta Concordancia (≥ 0.75): El análisis revela que {candidate_name} tiene una excelente adecuación con las funciones "
            f"del cargo de {cargo} y el perfil buscado. La experiencia detallada en su hoja de vida está estrechamente alineada con "
            f"las responsabilidades y competencias requeridas para este rol crucial en la prevalencia del Capítulo. La alta concordancia "
            f"indica que {candidate_name} está bien preparado para asumir este cargo y contribuir significativamente al éxito y la misión "
            f"del Capítulo. Se recomienda proceder con el proceso de selección y considerar a {candidate_name} como una opción sólida para el "
            f"cargo.")
    doc.add_paragraph(
        f"Este análisis es generado debido a que es crucial tomar medidas estratégicas para garantizar que los candidatos estén bien preparados "
        f"para el rol de {cargo}. Los aspirantes con alta concordancia deben ser considerados seriamente para el cargo, ya que están en una "
        f"posición ideal para contribuir efectivamente al liderazgo de la Junta Directiva Capitular.")

    # Agregar agradecimiento
    doc.add_paragraph(f"Muchas gracias {candidate_name} por tu interés en convertirte en {cargo}, ¡éxitos en tu proceso!")

    # Guardar archivo de reporte
    report_filename = f"reporte_{candidate_name}_{cargo}.docx"
    doc.save(report_filename)

    return report_filename

    if uploaded_zip and cargo and api_key:
        # Cargar archivos de funciones y perfil
        docs = load_job_docs(cargo)
        func_text = Document(docs["funciones"]).text
        profile_text = Document(docs["perfil"]).text

        # Extraer archivos .docx del ZIP
        docx_files = extract_docx_from_zip(uploaded_zip)

        # Procesar documentos y analizar
        experience_text = ""
        for docx_file in docx_files:
            doc = Document(docx_file)
            experience_text += extract_experience_aneiap(doc) + "\n"

        api_key = st.text_input("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")

        # Generar el reporte
        report_filename = generate_report(experience_text, func_text, profile_text, cargo, "Candidato", api_key)
        st.success(f"Reporte generado: {report_filename}")
        st.download_button("Descargar reporte", report_filename)
        
# Interfaz para subir el archivo .zip
uploaded_zip = st.file_uploader("Suba un archivo .zip con hojas de vida ANEIAP", type="zip")

# Función para manejar la carga y análisis
candidate_name = st.text_input("Nombre del candidato:")
cargo = st.selectbox("Seleccione el cargo que desea evaluar:", ["PC", "DCA", "DCC", "DCD", "DCF", "DCM", "CCP", "IC"])

if st.button("Evaluar"):
    if not candidate_name or not cargo or not uploaded_zip:
        st.error("Por favor, llena todos los campos y carga tu hoja de vida.")

if uploaded_zip:
    handle_uploaded_files(uploaded_zip)
